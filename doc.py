from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def agregar_titulo_seccion(doc, texto, nivel=1):
    heading = doc.add_heading(texto, level=nivel)
    return heading

def agregar_tabla_modelo(doc, titulo, columnas):
    doc.add_heading(titulo, level=3)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    for i, text in enumerate(['Campo', 'Tipo de Dato (PostgreSQL/Django)', 'Reglas / Restricciones']):
        run = hdr_cells[i].paragraphs[0].add_run(text)
        run.bold = True
    
    for campo, tipo, regla in columnas:
        row_cells = tabla.add_row().cells
        row_cells[0].text = campo
        row_cells[1].text = tipo
        row_cells[2].text = regla
    doc.add_paragraph('\n')

def agregar_historia(doc, us_data):
    doc.add_heading(f"{us_data['id']}: {us_data['titulo']}", level=3)
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = 'Table Grid'
    
    items = [
        ("Descripción", us_data['descripcion']),
        ("Criterios de Aceptación", us_data['criterios']),
        ("Reglas de Negocio / Backend", us_data['reglas']),
        ("Endpoints Asociados", us_data['endpoints']),
        ("Tareas Técnicas (Dev)", us_data['tareas'])
    ]
    
    for i, (campo, descripcion) in enumerate(items):
        celda_izq = tabla.cell(i, 0)
        texto_izq = celda_izq.paragraphs[0].add_run(campo)
        texto_izq.bold = True
        # Ancho aproximado
        celda_izq.width = Pt(150) 
        celda_der = tabla.cell(i, 1)
        celda_der.text = descripcion
    doc.add_paragraph('\n')

def generar_documento_maestro():
    doc = Document()
    titulo_principal = doc.add_heading('Software Requirements Specification (SRS)', 0)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Sistema de Simulacros Pruebas Saber Pro - Universidad Surcolombiana\nVersión: 1.0 (Definitiva)\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==========================================
    # SECCIÓN 1: ARQUITECTURA E INFRAESTRUCTURA
    # ==========================================
    agregar_titulo_seccion(doc, '1. Notas Arquitectónicas e Infraestructura')
    doc.add_paragraph('• Seguridad (Auth): Implementación estricta de JWT (JSON Web Tokens) con access_token (15 min) y refresh_token (24h) almacenados de forma segura en memoria o HttpOnly cookies en el frontend.')
    doc.add_paragraph('• Almacenamiento Multimedia (CDN): Para evitar colapso del servidor, las imágenes NUNCA se guardan en el disco duro local. Se requiere integración con AWS S3, Google Cloud Storage o equivalente.')
    doc.add_paragraph('• Procesamiento Asíncrono: Uso de Celery + Redis para la exportación de sábanas de datos pesadas (Excel) y generación de opciones con IA.')
    doc.add_paragraph('• Resiliencia Offline: Uso de localStorage/IndexedDB en React para garantizar que si el estudiante pierde internet, pueda seguir respondiendo y sincronizar al reconectar.\n')

    # ==========================================
    # SECCIÓN 2: MODELO DE DATOS (BASE DE DATOS)
    # ==========================================
    agregar_titulo_seccion(doc, '2. Diccionario de Datos Global (PostgreSQL)')

    agregar_tabla_modelo(doc, 'Tabla: Usuario', [
        ('id', 'UUID', 'Primary Key autogenerada.'),
        ('documento', 'VARCHAR(20)', 'Unique. Default initial password.'),
        ('tipo_documento', 'VARCHAR(5)', 'Choices: CC, TI, CE, PEP.'),
        ('nombres / apellidos', 'VARCHAR(100)', 'No Null.'),
        ('correo_institucional', 'VARCHAR(150)', 'Unique. Username field for login (must contain @usco.edu.co).'),
        ('programa', 'VARCHAR(150)', 'No Null.'),
        ('es_primer_ingreso', 'Boolean', 'Default: True. Fuerza cambio de clave.'),
        ('is_active', 'Boolean', 'Default: True.')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Pregunta', [
        ('id', 'UUID', 'Primary Key.'),
        ('modulo_id / categoria_id / competencia_id', 'ForeignKey', 'Relación jerárquica obligatoria.'),
        ('nivel_dificultad', 'VARCHAR(15)', 'Fácil, Medio, Difícil.'),
        ('contexto_texto / contexto_imagen', 'TEXT / ImageField', 'Nullable. Texto base o gráfica previa al enunciado.'),
        ('enunciado', 'TEXT', 'No Null. Soporta HTML y LaTeX.'),
        ('limite_palabras / rubrica', 'Integer / TEXT', 'Solo para módulo Comunicación Escrita.'),
        ('estado', 'VARCHAR(15)', 'Borrador, Publicada, Archivada (Borrado lógico).'),
        ('version_original_id', 'ForeignKey (Self)', 'Nullable. Apunta a la pregunta original si fue versionada por edición.')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Opcion_Respuesta', [
        ('id', 'UUID', 'Primary Key.'),
        ('pregunta_id', 'ForeignKey', 'Relación CASCADE hacia tabla Pregunta.'),
        ('texto / imagen', 'TEXT / ImageField', 'Contenido de la opción. Soporta LaTeX.'),
        ('es_correcta', 'Boolean', 'No Null. Excluida estrictamente del payload público (Serializer).')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Plantilla_Examen', [
        ('id', 'UUID', 'Primary Key.'),
        ('tiempo_minutos', 'Integer', 'Nullable (Indefinido si es null/0).'),
        ('fecha_inicio / fecha_fin', 'DateTime', 'Ventana estricta de disponibilidad.'),
        ('mostrar_resultados_inmediatos', 'Boolean', 'Política antifraude.'),
        ('es_simulacro_oficial', 'Boolean', 'Fuerza reglas fijas del ICFES.')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Regla_Examen', [
        ('id', 'UUID', 'Primary Key.'),
        ('examen_id', 'ForeignKey', 'Hacia Plantilla_Examen.'),
        ('modulo_id / categoria_id', 'ForeignKey', 'De dónde extraer.'),
        ('cantidad_preguntas', 'Integer', 'Número de ítems a extraer.'),
        ('nivel_dificultad', 'VARCHAR(15)', 'Default: Balanceada (33/33/34).')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Intento_Examen (El Estudiante)', [
        ('id', 'UUID', 'Primary Key.'),
        ('estudiante_id / plantilla_examen_id', 'ForeignKey', 'Relaciones obligatorias.'),
        ('fecha_inicio / fecha_finalizacion', 'DateTime', 'Control de tiempo regido por servidor.'),
        ('estado', 'VARCHAR(30)', 'En Progreso, Pendiente Calificacion, Finalizado.')
    ])

    agregar_tabla_modelo(doc, 'Tabla: Respuesta_Estudiante', [
        ('id', 'UUID', 'Primary Key.'),
        ('intento_id / pregunta_id', 'ForeignKey', 'Registro de pregunta extraída.'),
        ('opcion_seleccionada_id', 'ForeignKey', 'Nullable. Si es null, pregunta en blanco.'),
        ('texto_respuesta_abierta', 'TEXT', 'Nullable. Ensayo escrito.'),
        ('puntaje_calificado', 'Decimal(5,2)', 'Autocalculado o manual (US-010).')
    ])

    doc.add_page_break()

    # ==========================================
    # SECCIÓN 3: HISTORIAS DE USUARIO (EPICS)
    # ==========================================
    agregar_titulo_seccion(doc, '3. Especificación de Historias de Usuario')

    historias = [
        {
            "id": "US-001", "titulo": "Inicio de Sesión y Autenticación",
            "descripcion": "Login seguro verificando dominio institucional y roles.",
            "criterios": "1. Validación Regex de @usco.edu.co. 2. Endpoint retorna JWT. 3. Manejo de error genérico ('Credenciales incorrectas') por seguridad.",
            "reglas": "Uso de djangorestframework-simplejwt.",
            "endpoints": "POST /api/auth/login/",
            "tareas": "- Configurar JWT.\n- Setup React Context para estado de sesión."
        },
        {
            "id": "US-002", "titulo": "Activación de Cuenta (Primer Ingreso)",
            "descripcion": "Cambio de contraseña obligatorio si es la primera vez que ingresa.",
            "criterios": "1. Bloqueo de navegación a otras rutas hasta cambiar clave. 2. Validación: Min 8 chars, 1 mayúscula, 1 número. No igual a documento.",
            "reglas": "Hash con Argon2 o PBKDF2. Actualiza es_primer_ingreso a False.",
            "endpoints": "POST /api/auth/activar/",
            "tareas": "- Middleware en React para forzar redirección a /activacion."
        },
        {
            "id": "US-003", "titulo": "Recuperación de Contraseña",
            "descripcion": "Solicitud de enlace de reseteo vía correo electrónico.",
            "criterios": "1. Token único caduca en 15 min. 2. Rate limiting (max 5 peticiones) para evitar spam.",
            "reglas": "Tokens de un solo uso (invalidar tras uso).",
            "endpoints": "POST /api/auth/recuperar/ \nPOST /api/auth/reset/",
            "tareas": "- Integrar SMTP (SendGrid/AWS SES).\n- Lógica de invalidación de token."
        },
        {
            "id": "US-004", "titulo": "Cierre de Sesión Seguro (Logout)",
            "descripcion": "Destrucción de tokens y limpieza local.",
            "criterios": "1. Backend añade refresh_token a Blacklist. 2. Frontend limpia localStorage, sessionStorage y Context.",
            "reglas": "Try/Finally en frontend para garantizar limpieza local incluso sin internet.",
            "endpoints": "POST /api/auth/logout/",
            "tareas": "- Habilitar app token_blacklist en Django."
        },
        {
            "id": "US-005", "titulo": "Perfil de Usuario",
            "descripcion": "Vista de datos académicos (solo lectura) y cambio voluntario de clave.",
            "criterios": "1. Datos sincronizados fijos. 2. Exigir 'Contraseña Actual' para realizar el cambio.",
            "reglas": "Validar clave antigua antes de aplicar hash a la nueva.",
            "endpoints": "GET /api/perfil/ \nPUT /api/perfil/cambiar-clave/",
            "tareas": "- Diseño UI del perfil."
        },
        {
            "id": "US-006", "titulo": "Gestión Individual y Carga Masiva de Usuarios",
            "descripcion": "CRUD de estudiantes por parte del admin.",
            "criterios": "1. Carga Excel (.xlsx) con filas atómicas (si falla una, reporta y sigue). 2. Tabla Admin para editar correo, desactivar acceso (retiro) o forzar reseteo de clave por bloqueo.",
            "reglas": "Uso estricto de transaction.atomic(). Clave inicial = Documento.",
            "endpoints": "POST /api/admin/usuarios/carga-masiva/ \nPUT /api/admin/usuarios/{id}/",
            "tareas": "- Implementar pandas/openpyxl.\n- Modal Drag & Drop con reporte de errores."
        },
        {
            "id": "US-013", "titulo": "Parametrización del Sistema",
            "descripcion": "Superadmin configura Módulos, Categorías y Competencias base.",
            "criterios": "1. CRUD para jerarquías (Módulo -> Categoría -> Competencia). 2. Alimenta los selectores del Banco de Preguntas.",
            "reglas": "No se pueden eliminar categorías con preguntas asociadas (solo Inactivar).",
            "endpoints": "CRUD en /api/admin/parametros/",
            "tareas": "- Seed inicial de la base de datos con los 5 módulos Saber Pro."
        },
        {
            "id": "US-007", "titulo": "Gestión Avanzada del Banco de Preguntas",
            "descripcion": "CRUD de preguntas con Contexto, LaTeX, IA y Reglas de Versionamiento.",
            "criterios": "1. Editor RichText (LaTeX). 2. Subida de imágenes. 3. Botón 'Sugerir Distractores con IA'. 4. Versionamiento: Si se edita una pregunta ya usada, crear clon (v2) y archivar v1. 5. Borrado lógico.",
            "reglas": "SEGURIDAD CRÍTICA: Payload JAMÁS incluye 'es_correcta' ni 'justificacion' hacia el estudiante. Las imágenes se comprimen a max 800px en backend.",
            "endpoints": "POST /api/admin/preguntas/ (Nested JSON) \nPOST /api/admin/ia/distractores/",
            "tareas": "- Configurar Nested Serializers.\n- Integrar API Gemini/LLM.\n- Optimización de imágenes (Pillow)."
        },
        {
            "id": "US-008", "titulo": "Ensamblaje de Plantillas de Exámenes",
            "descripcion": "Configuración de reglas automáticas de extracción para los simulacros.",
            "criterios": "1. Fechas y público objetivo (Programas). 2. Políticas de resultados. 3. Constructor de Reglas (Bloques). 4. Dificultad Balanceada (33% fácil, 33% medio, 34% difícil). 5. Validación de stock de preguntas en BD.",
            "reglas": "Lógica backend para no repetir preguntas al estudiante entre reglas del mismo módulo. Bloqueo de edición si ya hay intentos iniciados.",
            "endpoints": "POST /api/admin/examenes/ \nGET /api/admin/preguntas/conteo/",
            "tareas": "- Lógica de cálculo 33/33/34 en Backend."
        },
        {
            "id": "US-009", "titulo": "Realización del Simulacro (Módulo Estudiante)",
            "descripcion": "Motor de examen con navegación por módulos y resiliencia offline.",
            "criterios": "1. Generación de Intento inmutable con timestamp. 2. Pantalla Lobby con estado de 5 módulos. 3. Dentro del módulo: 1 pregunta por pantalla y cuadrícula de navegación. 4. Bloqueo de módulo al darle 'Terminar Módulo'. 5. Autoguardado Asíncrono.",
            "reglas": "Si falla la red, guardar en IndexedDB y sincronizar al volver online. Temporizador estrictamente validado contra servidor, no reloj local.",
            "endpoints": "POST /api/examenes/{id}/iniciar/ \nPOST /api/examenes/intentos/{id}/responder/",
            "tareas": "- Lógica de cola offline/online en React.\n- Generador aleatorio de examen Anti-Duplicados."
        },
        {
            "id": "US-010", "titulo": "Calificación Manual (Comunicación Escrita)",
            "descripcion": "Evaluación de ensayos por docentes (Split-Screen).",
            "criterios": "1. Bandeja de pendientes. 2. Split-Screen (Ensayo vs Rúbrica/Nota). 3. Candado de concurrencia (Pessimistic Locking) si otro profe ya lo abrió. 4. Autocalificación en 0 si ensayo está en blanco.",
            "reglas": "Al guardar, suma nota manual a autocalificación Múltiple y pasa a Finalizado.",
            "endpoints": "POST /api/admin/evaluaciones/{id}/calificar/",
            "tareas": "- select_for_update() en Django para bloquear fila concurrente.\n- Celery task para ensayos vacíos."
        },
        {
            "id": "US-011", "titulo": "Dashboard de Resultados (Estudiante)",
            "descripcion": "Retroalimentación de notas normalizadas y revisión por ítem.",
            "criterios": "1. Bloqueo por política de fecha de cierre. 2. Estado 'Calificación Parcial' si falta ensayo. 3. Gráficos de barras/radar. 4. Detalle: Respuesta marcada vs correcta + Justificación.",
            "reglas": "Fórmula de Normalización: (Aciertos / Total) * 300. Optimización de query obligatoria para evitar Timeout.",
            "endpoints": "GET /api/estudiantes/mis-resultados/{id}/resumen/",
            "tareas": "- Implementar select_related en Django.\n- Integrar Recharts/Chart.js en frontend."
        },
        {
            "id": "US-012", "titulo": "Inteligencia Académica y Dashboard (Admin)",
            "descripcion": "Analíticas globales para directivas y profesores.",
            "criterios": "1. Filtros (Programa, Sede). 2. KPIs Globales. 3. Top 10 Estudiantes. 4. Análisis de Ítems (Preguntas con mayor % de error). 5. Exportación a Excel (.xlsx).",
            "reglas": "Uso estricto de caché de Django (Redis) para evitar colapsar la DB con promedios. Exportaciones pesadas deben ser asíncronas.",
            "endpoints": "GET /api/admin/analiticas/kpis/ \nGET /api/admin/analiticas/exportar/",
            "tareas": "- Configurar Redis cache.\n- Pandas/Openpyxl para generar Excel binario."
        }
    ]

    for us in historias:
        agregar_historia(doc, us)

    nombre_archivo = 'SRS_SaberPro_USCO_Definitivo_Final.docx'
    doc.save(nombre_archivo)
    print(f"¡DOCUMENTO MAESTRO GENERADO! Archivo '{nombre_archivo}' guardado exitosamente. Todo tu trabajo está aquí.")

if __name__ == "__main__":
    generar_documento_maestro()